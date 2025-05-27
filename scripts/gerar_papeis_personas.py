import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime

# ========================================
# 📌 CONFIGURAÇÕES INSTITUCIONAIS
# ========================================
cliente = "Up Séries"
autores = ["Arimatéia Júnior"]
empresa_exec = "StartCode Cloud Tecnologia"
logo_path = "dados/logo-startcodecloud.png"
data_hoje = datetime.now().strftime("%d/%m/%Y")

# ========================================
# 📂 ARQUIVOS DE ENTRADA E SAÍDA
# ========================================
entrada = "dados/02_papeis_personas.xlsx"
saida = "relatorios/Relatorio_Papeis_Personas.docx"
titulo = "Relatório de Papéis e Personas FinOps"

# ========================================
# 📥 LEITURA DOS DADOS
# ========================================
df = pd.read_excel(entrada)
df.columns = [col.strip() for col in df.columns]

# ========================================
# 📝 CRIAÇÃO DO DOCUMENTO
# ========================================
doc = Document()

# 🔰 CABEÇALHO COM TÍTULO E LOGO
table = doc.add_table(rows=1, cols=2)
table.allow_autofit = True
cell_titulo = table.cell(0, 0)
cell_logo = table.cell(0, 1)

p_titulo = cell_titulo.paragraphs[0]
run = p_titulo.add_run(titulo)
run.bold = True
run.font.size = doc.styles['Normal'].font.size
p_titulo.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
cell_titulo.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

try:
    run_logo = cell_logo.paragraphs[0].add_run()
    run_logo.add_picture(logo_path, width=Inches(1.2))
    cell_logo.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
except:
    print("⚠️ Logo não encontrada. Pulando...")

doc.add_paragraph()

# 🧾 TABELA DE METADADOS
meta_table = doc.add_table(rows=2, cols=2)
meta_table.style = 'Table Grid'
meta_table.autofit = True

meta_table.cell(0, 0).text = f"📛 Cliente: {cliente}"
meta_table.cell(0, 1).text = f"🏢 Empresa Executora: {empresa_exec}"
meta_table.cell(1, 0).text = f"🧑‍💼 Autores: {', '.join(autores)}"
meta_table.cell(1, 1).text = f"📅 Data do Relatório: {data_hoje}"

doc.add_paragraph()

# 📋 TABELA DE PAPÉIS E RESPONSABILIDADES
doc.add_heading("Papéis e Responsabilidades FinOps", level=1)
tabela = doc.add_table(rows=1, cols=len(df.columns))
tabela.style = 'Table Grid'
hdr = tabela.rows[0].cells
for i, col in enumerate(df.columns):
    hdr[i].text = col

for _, row in df.iterrows():
    row_cells = tabela.add_row().cells
    for i, item in enumerate(row):
        row_cells[i].text = str(item)

doc.add_paragraph()

# 📊 GRÁFICO DE PARTICIPAÇÃO POR PERSONA
doc.add_heading("Gráfico: Participação por Persona", level=1)

if "Responsável" in df.columns:
    # Contar quantos papéis cada responsável ocupa
    contagem = df["Responsável"].value_counts()
    percentuais = contagem / contagem.sum() * 100

    # Juntar descrições das funções por responsável
    funcoes_agrupadas = df.groupby("Responsável")["Descrição da Função"].apply(
        lambda x: "\n".join(f"• {i}" for i in x)
    )

    # Criar gráfico
    plt.figure(figsize=(8, 5))
    bars = plt.bar(percentuais.index, percentuais.values, color="#4C9AFF")
    plt.title("Distribuição de Responsabilidades FinOps por Pessoa")
    plt.ylabel("Percentual (%)")
    plt.ylim(0, 100)
    plt.xticks(rotation=0)
    plt.grid(axis='y', linestyle='--', alpha=0.3)

    # Rótulo com % e funções
    for bar, nome in zip(bars, percentuais.index):
        altura = bar.get_height()
        texto_funcoes = funcoes_agrupadas.get(nome, "")
        texto_rotulo = f"{altura:.0f}%\n{texto_funcoes}"
        plt.text(
            bar.get_x() + bar.get_width() / 2,
            altura + 1,
            texto_rotulo,
            ha='center',
            fontsize=8
        )

    plt.tight_layout()
    grafico_path = "grafico_papeis_personas.png"
    plt.savefig(grafico_path)
    plt.close()

    doc.add_picture(grafico_path, width=Inches(5))
else:
    doc.add_paragraph("Nenhuma coluna 'Responsável' encontrada para gerar gráfico.")

# 💾 SALVAR DOCUMENTO
doc.save(saida)
print("✅ Relatório gerado com sucesso:", saida)
