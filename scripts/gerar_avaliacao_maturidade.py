
import pandas as pd
import matplotlib.pyplot as plt
import numpy as np
from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
import os

# ========================================
# 📌 CONFIGURAÇÕES INSTITUCIONAIS DO RELATÓRIO
# ========================================
cliente = "Up Séries"
autores = ["Arimatéia Júnior"]
empresa_exec = "StartCode Cloud Tecnologia"
logo_path = "dados/logo-startcodecloud.png"
data_hoje = datetime.now().strftime("%d/%m/%Y")

# ========================================
# 📂 ARQUIVOS DE ENTRADA/SAÍDA
# ========================================
entrada = "dados/01_avaliacao_maturidade.xlsx"
saida = "relatorios/Relatorio_Avaliacao_Maturidade.docx"
titulo = "Relatório de Avaliação de Maturidade FinOps"

# ========================================
# ✅ DOMÍNIOS VÁLIDOS A SEREM CONSIDERADOS
# ========================================
dominios_validos = [
    "Entender Uso e Custo",
    "Quantificar Valor de Negócio",
    "Otimizar Uso e Custo",
    "Governança e Operação"
]

# ========================================
# 📥 LEITURA E LIMPEZA DOS DADOS
# ========================================
df = pd.read_excel(entrada, usecols="A:D")
df.columns = [col.strip() for col in df.columns]
df = df[df["Domínio"].isin(dominios_validos)]
df = df[pd.to_numeric(df["Maturidade"], errors='coerce').notna()]
df["Maturidade"] = df["Maturidade"].astype(int)

# ========================================
# 📊 AGRUPAMENTO POR DOMÍNIO E CLASSIFICAÇÃO
# ========================================
agrupado = df.groupby("Domínio")["Maturidade"].mean().reset_index()
agrupado["Classificação"] = agrupado["Maturidade"].apply(
    lambda x: "Rastejar" if x <= 1.4 else "Andar" if x <= 2.4 else "Correr"
)

# ========================================
# 📉 GRÁFICO DE BARRAS
# ========================================
cores = {
    "Rastejar": "#FF4D4D",
    "Andar": "#FFC107",
    "Correr": "#4CAF50"
}
agrupado["Cor"] = agrupado["Classificação"].map(cores)

plt.figure(figsize=(7, 4))
bars = plt.bar(
    agrupado["Domínio"],
    agrupado["Maturidade"],
    color=agrupado["Cor"]
)
plt.title("Maturidade por Domínio (Classificada)")
plt.ylabel("Nível (1 = Rastejar, 3 = Correr)")
plt.ylim(0, 3.2)
plt.xticks(rotation=45, ha='right')
plt.grid(axis='y', linestyle='--', alpha=0.3)

for bar, texto in zip(bars, agrupado["Classificação"]):
    plt.text(
        bar.get_x() + bar.get_width() / 2,
        bar.get_height() + 0.1,
        texto,
        ha='center',
        fontsize=9,
        weight='bold'
    )

plt.tight_layout()
grafico_barras = "grafico_maturidade.png"
plt.savefig(grafico_barras)
plt.close()

# ========================================
# 📉 GRÁFICO DE RADAR
# ========================================
dominios = agrupado["Domínio"].tolist()
valores = agrupado["Maturidade"].tolist()
valores += valores[:1]
angles = np.linspace(0, 2 * np.pi, len(dominios), endpoint=False).tolist()
angles += angles[:1]

fig, ax = plt.subplots(figsize=(6, 6), subplot_kw=dict(polar=True))
ax.plot(angles, valores, color='blue', linewidth=2, label='Média')
ax.fill(angles, valores, color='skyblue', alpha=0.3)
ax.set_yticklabels([])
ax.set_xticks(angles[:-1])
ax.set_xticklabels(dominios)
plt.title("Maturidade por Domínio", size=15, weight='bold', pad=20)
plt.legend(loc='upper right', bbox_to_anchor=(1.2, 1.1))
grafico_radar = "grafico_radar_maturidade.png"
plt.savefig(grafico_radar, bbox_inches='tight')
plt.close()

# ========================================
# 📝 CRIAÇÃO DO DOCUMENTO WORD
# ========================================
doc = Document()
table = doc.add_table(rows=1, cols=2)
table.allow_autofit = True
cell_titulo = table.cell(0, 0)
cell_logo = table.cell(0, 1)

p_titulo = cell_titulo.paragraphs[0]
run = p_titulo.add_run(titulo)
run.bold = True
run.font.size = doc.styles['Normal'].font.size
p_titulo.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
cell_titulo.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

try:
    run_logo = cell_logo.paragraphs[0].add_run()
    run_logo.add_picture(logo_path, width=Inches(1.2))
    cell_logo.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
except:
    print("⚠️ Logo não encontrada. Pulando...")

doc.add_paragraph()

meta_table = doc.add_table(rows=2, cols=2)
meta_table.style = 'Table Grid'
meta_table.cell(0, 0).text = f"📛 Cliente: {cliente}"
meta_table.cell(0, 1).text = f"🏢 Empresa Executora: {empresa_exec}"
meta_table.cell(1, 0).text = f"🧑‍💼 Autores: {', '.join(autores)}"
meta_table.cell(1, 1).text = f"📅 Data do Relatório: {data_hoje}"

doc.add_paragraph()

# =====================================
# 📊 TABELA DE AVALIAÇÃO
# =====================================
doc.add_heading("Tabela de Avaliação", level=1)
tabela = doc.add_table(rows=1, cols=len(df.columns))
tabela.style = 'Table Grid'
hdr = tabela.rows[0].cells
for i, col in enumerate(df.columns):
    hdr[i].text = col

for _, row in df.iterrows():
    row_cells = tabela.add_row().cells
    for i, item in enumerate(row):
        row_cells[i].text = str(item)

# Inserir gráficos
doc.add_heading("Gráfico: Classificação por Domínio", level=1)
doc.add_picture(grafico_barras, width=Inches(5))

doc.add_heading("Gráfico Radar de Maturidade", level=1)
doc.add_picture(grafico_radar, width=Inches(5.5))

# =====================================
# 💾 SALVAR O RELATÓRIO
# =====================================
doc.save(saida)
print("✅ Relatório gerado com sucesso:", saida)
