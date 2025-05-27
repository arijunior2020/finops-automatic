import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime

# ========================================
# 📌 CONFIGURAÇÕES INSTITUCIONAIS
# ========================================
cliente = "Antagonista"
autores = ["Arimatéia Júnior", "Yuri Lopes"]
empresa_exec = "InCloud tecnologia"
logo_path = "dados/logo_incloud.png"
data_hoje = datetime.now().strftime("%d/%m/%Y")

# ========================================
# 📂 ARQUIVOS DE ENTRADA E SAÍDA
# ========================================
entrada = "dados/04_execucao_piloto.xlsx"
saida = "relatorios/Relatorio_Execucao_Piloto.docx"
titulo = "Relatório de Execução do Piloto FinOps"

# ========================================
# 📥 LEITURA DO ARQUIVO
# ========================================
df = pd.read_excel(entrada)

# ========================================
# 📝 CRIAÇÃO DO DOCUMENTO WORD
# ========================================
doc = Document()

# ================================
# 🔰 CABEÇALHO COM TÍTULO + LOGO
# ================================
table = doc.add_table(rows=1, cols=2)
table.allow_autofit = True
cell_titulo = table.cell(0, 0)
cell_logo = table.cell(0, 1)

p_titulo = cell_titulo.paragraphs[0]
run = p_titulo.add_run(titulo)
run.bold = True
run.font.size = doc.styles['Normal'].font.size
p_titulo.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
cell_titulo.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

try:
    run_logo = cell_logo.paragraphs[0].add_run()
    run_logo.add_picture(logo_path, width=Inches(1.2))
    cell_logo.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
except:
    print("⚠️ Logo não encontrada. Pulando...")

doc.add_paragraph()  # Espaço após cabeçalho

# =====================================
# 🧾 TABELA DE METADADOS
# =====================================
meta_table = doc.add_table(rows=2, cols=2)
meta_table.style = 'Table Grid'
meta_table.autofit = True

meta_table.cell(0, 0).text = f"📛 Cliente: {cliente}"
meta_table.cell(0, 1).text = f"🏢 Empresa Executora: {empresa_exec}"
meta_table.cell(1, 0).text = f"🧑‍💼 Autores: {', '.join(autores)}"
meta_table.cell(1, 1).text = f"📅 Data do Relatório: {data_hoje}"

doc.add_paragraph()  # Espaço após metadados

# =====================================
# 📋 TABELA DE DADOS DO PILOTO
# =====================================
doc.add_heading("Projetos Avaliados no Piloto", level=1)
tabela = doc.add_table(rows=1, cols=len(df.columns))
tabela.style = 'Table Grid'
hdr = tabela.rows[0].cells
for i, col in enumerate(df.columns):
    hdr[i].text = col

for _, row in df.iterrows():
    row_cells = tabela.add_row().cells
    for i, item in enumerate(row):
        row_cells[i].text = str(item)

# =====================================
# 📈 GRÁFICO DE RESULTADO DO PILOTO
# =====================================
doc.add_heading("Gráfico: Resultado por Projeto (%)", level=1)

plt.figure(figsize=(7, 4))
bars = plt.bar(
    df["Projeto"],
    df["Resultado (%)"],
    color="#4C9AFF"
)

plt.title("Resultados do Piloto (%)")
plt.ylabel("Percentual")
plt.ylim(0, 110)
plt.xticks(rotation=0)
plt.grid(axis='y', linestyle='--', alpha=0.3)

# Adiciona os rótulos nas barras
for bar, val in zip(bars, df["Resultado (%)"]):
    plt.text(
        bar.get_x() + bar.get_width() / 2,
        bar.get_height() + 2,
        f"{val}%",
        ha='center',
        fontsize=9,
        weight='bold'
    )

plt.tight_layout()
grafico_path = "grafico_execucao_piloto.png"
plt.savefig(grafico_path)
plt.close()

doc.add_picture(grafico_path, width=Inches(5))

# =====================================
# 💾 SALVAR RELATÓRIO
# =====================================
doc.save(saida)
print("✅ Relatório gerado com sucesso:", saida)
