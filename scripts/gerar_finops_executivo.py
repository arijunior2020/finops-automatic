
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime
import os

# 📁 Configurações iniciais
entrada = "dados/06_finops_executivo.xlsx"
saida = "relatorios/Relatorio_FinOps_Executivo.docx"
grafico_categorias = "grafico_categorias.png"
logo_path = "dados/logo-startcodecloud.png"
cliente = "UP Séries"
autores = ["Arimatéia Júnior"]
empresa_exec = "StartCode Cloud Tecnologia"
data_hoje = datetime.now().strftime("%d/%m/%Y")

# 📥 Leitura dos dados preenchidos
df = pd.read_excel(entrada)
df.columns = [col.strip() for col in df.columns]

# 📊 Gerar gráfico de distribuição por categoria
df["Categoria"] = df["Categoria"].str.strip()
contagem = df["Categoria"].value_counts().reset_index()
contagem.columns = ["Categoria", "Total"]

plt.figure(figsize=(8, 4))
bars = plt.bar(contagem["Categoria"], contagem["Total"], color="skyblue")
plt.title("Distribuição de Itens Avaliados por Categoria")
plt.ylabel("Quantidade de Itens")
plt.xticks(rotation=45)
plt.tight_layout()
plt.savefig(grafico_categorias)
plt.close()

# 📄 Criar documento Word
doc = Document()

# 🔰 Cabeçalho com título e logo
table = doc.add_table(rows=1, cols=2)
cell_title, cell_logo = table.cell(0, 0), table.cell(0, 1)

p_title = cell_title.paragraphs[0]
p_title.add_run("Relatório Executivo FinOps AWS").bold = True
p_title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
cell_title.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

if logo_path and os.path.exists(logo_path):
    run_logo = cell_logo.paragraphs[0].add_run()
    run_logo.add_picture(logo_path, width=Inches(1.2))
    cell_logo.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

doc.add_paragraph()

# 🧾 Metadados institucionais
meta = doc.add_table(rows=2, cols=2)
meta.style = 'Table Grid'
meta.cell(0, 0).text = f"📛 Cliente: {cliente}"
meta.cell(0, 1).text = f"🏢 Empresa Executora: {empresa_exec}"
meta.cell(1, 0).text = f"🧑‍💼 Autores: {', '.join(autores)}"
meta.cell(1, 1).text = f"📅 Data do Relatório: {data_hoje}"
doc.add_paragraph()

# 📈 Inserir gráfico
doc.add_heading("Gráfico de Distribuição por Categoria", level=1)
if os.path.exists(grafico_categorias):
    doc.add_picture(grafico_categorias, width=Inches(5.5))
    doc.add_paragraph()

# 📝 Tabela geral
doc.add_heading("Tabela Consolidada de Avaliação", level=1)
tabela = doc.add_table(rows=1, cols=len(df.columns))
tabela.style = 'Table Grid'
hdr = tabela.rows[0].cells
for i, col in enumerate(df.columns):
    hdr[i].text = col

for _, row in df.iterrows():
    row_cells = tabela.add_row().cells
    for i, item in enumerate(row):
        row_cells[i].text = str(item)

doc.add_paragraph()

# 🧭 Seções por Categoria
for categoria in df["Categoria"].unique():
    doc.add_heading(categoria, level=1)
    subset = df[df["Categoria"] == categoria]

    for _, row in subset.iterrows():
        doc.add_heading(row["Domínio/Tema"], level=2)
        doc.add_paragraph(f"📌 *Item Avaliado:* {row['Descrição do Item Avaliado']}")
        doc.add_paragraph(f"📝 *Resposta/Avaliação:* {row['Resposta/Avaliação']}")
        doc.add_paragraph("")

# ✅ Etapas de Plano de Ação
doc.add_heading("Plano de Ação Consolidado", level=1)
acoes = df[df["Categoria"] == "Recomendações"]
for _, row in acoes.iterrows():
    doc.add_paragraph(f"🔹 {row['Domínio/Tema']}: {row['Resposta/Avaliação']}")

# 💾 Salvar o relatório
doc.save(saida)
print("✅ Relatório gerado com sucesso:", saida)
