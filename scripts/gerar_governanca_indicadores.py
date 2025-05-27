import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime

# ========================================
# 📌 CONFIGURAÇÕES INSTITUCIONAIS
# ========================================
cliente = "Up Séries"
autores = ["Arimatéia Júnior"]
empresa_exec = "StartCode Cloud Tecnologia"
logo_path = "dados/logo-startcodecloud.png"
data_hoje = datetime.now().strftime("%d/%m/%Y")

# ========================================
# 📂 ENTRADA E SAÍDA
# ========================================
entrada = "dados/05_governanca_indicadores.xlsx"
saida = "relatorios/Relatorio_Governanca_Indicadores.docx"
titulo = "Relatório de Governança e Indicadores FinOps"

# ========================================
# 📥 LEITURA DOS DADOS
# ========================================
df = pd.read_excel(entrada)

# ========================================
# 📝 CRIAÇÃO DO DOCUMENTO
# ========================================
doc = Document()

# ================================
# 🔰 CABEÇALHO COM TÍTULO E LOGO
# ================================
table = doc.add_table(rows=1, cols=2)
table.allow_autofit = True
cell_titulo = table.cell(0, 0)
cell_logo = table.cell(0, 1)

p_titulo = cell_titulo.paragraphs[0]
run = p_titulo.add_run(titulo)
run.bold = True
run.font.size = doc.styles['Normal'].font.size
p_titulo.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
cell_titulo.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

try:
    run_logo = cell_logo.paragraphs[0].add_run()
    run_logo.add_picture(logo_path, width=Inches(1.2))
    cell_logo.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
except:
    print("⚠️ Logo não encontrada. Pulando...")

doc.add_paragraph()

# ========================================
# 🧾 TABELA DE METADADOS
# ========================================
meta_table = doc.add_table(rows=2, cols=2)
meta_table.style = 'Table Grid'
meta_table.autofit = True
meta_table.cell(0, 0).text = f"📛 Cliente: {cliente}"
meta_table.cell(0, 1).text = f"🏢 Empresa Executora: {empresa_exec}"
meta_table.cell(1, 0).text = f"🧑‍💼 Autores: {', '.join(autores)}"
meta_table.cell(1, 1).text = f"📅 Data do Relatório: {data_hoje}"

doc.add_paragraph()

# ========================================
# 📋 TABELA DE INDICADORES
# ========================================
doc.add_heading("Indicadores de Governança", level=1)
tabela = doc.add_table(rows=1, cols=len(df.columns))
tabela.style = 'Table Grid'
hdr = tabela.rows[0].cells
for i, col in enumerate(df.columns):
    hdr[i].text = col

for _, row in df.iterrows():
    row_cells = tabela.add_row().cells
    for i, item in enumerate(row):
        row_cells[i].text = str(item)

# ========================================
# 📊 GRÁFICO DE COMPARAÇÃO: META VS RESULTADO
# ========================================
doc.add_heading("📉 Gráfico: Meta x Resultado dos Indicadores", level=1)

indicadores = df["Indicador"]
metas = df["Meta (%)"]
resultados = df["Resultado (%)"]

x = range(len(indicadores))
bar_width = 0.35

plt.figure(figsize=(8, 4))
plt.bar(x, metas, width=bar_width, label="Meta", color="#AAAAAA")
plt.bar([i + bar_width for i in x], resultados, width=bar_width, label="Resultado", color="#4C9AFF")

plt.xticks([i + bar_width / 2 for i in x], indicadores, rotation=0)
plt.title("Meta vs Resultado por Indicador")
plt.ylabel("Percentual (%)")
plt.ylim(0, 110)
plt.legend()
plt.grid(axis='y', linestyle='--', alpha=0.3)

for i, (meta, resultado) in enumerate(zip(metas, resultados)):
    plt.text(i, meta + 1, f"{meta}%", ha='center', fontsize=8)
    plt.text(i + bar_width, resultado + 1, f"{resultado}%", ha='center', fontsize=8)

plt.tight_layout()
grafico_path = "grafico_governanca_indicadores.png"
plt.savefig(grafico_path)
plt.close()

doc.add_picture(grafico_path, width=Inches(5))

# ========================================
# 💾 SALVAR O DOCUMENTO
# ========================================
doc.save(saida)
print("✅ Relatório gerado com sucesso:", saida)
